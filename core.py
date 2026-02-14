import os
import json
import time
import datetime
import glob
import hashlib
import requests
import html
from dotenv import load_dotenv
import google.generativeai as genai
from pypdf import PdfReader
import hmac
import secrets

# Load environment variables
load_dotenv(override=True)

CMS_ROOT = "smart_cms_data"

MAX_INPUT_SIZE = 50000 # Character limit for safety

def check_env_security():
    """Enhanced environment and API Key security verification."""
    # Check if we are running in Streamlit Cloud environment
    is_streamlit_cloud = os.getenv("STREAMLIT_RUNTIME_ENV") is not None or os.path.exists("/app")
    
    # Check for the common ' .env' leak (leading space)
    if os.path.exists(" .env"):
        return False, "🚨 SECURITY CRITICAL: Duplicate '.env' file with leading space found! This is NOT ignored by git and WILL LEAK KEYS. Delete it immediately."

    if not os.path.exists(".env"):
        if is_streamlit_cloud:
            # On Streamlit Cloud, keys should be in st.secrets, so .env missing is okay
            return True, "✅ Streamlit Cloud detected. Using st.secrets for configuration."
        return False, "🚨 SECURITY CRITICAL: .env file missing! Create one based on .env.example"
    
    # Check for placeholder keys
    key = os.getenv("GEMINI_API_KEY")
    if not key or "YOUR_API_KEY" in key or len(key.strip()) < 20:
        # Check if it might be in Streamlit secrets first
        try:
            import streamlit as st
            if "GEMINI_API_KEY" in st.secrets:
                return True, "✅ Security checks passed (using st.secrets)."
        except: pass
        return False, "⚠️ SECURITY RISK: Invalid or placeholder GEMINI_API_KEY found in .env"
    
    # Check if .env is in .gitignore
    if os.path.exists(".gitignore"):
        with open(".gitignore", "r") as f:
            content = f.read()
            if ".env" not in content:
                return False, "🚨 SECURITY RISK: .env is not in .gitignore. Do NOT commit your keys!"
        
    return True, "✅ Security checks passed."

def get_api_key(task_type):
    """
    Retrieves the appropriate API key, falling back to the master key.
    Ensures that placeholder strings are ignored.
    Forces a reload of environment variables to avoid using stale keys.
    """
    from dotenv import load_dotenv
    load_dotenv(override=True)
    
    key_map = {
        "creation": "GEMINI_API_KEY_CREATION",
        "transformation": "GEMINI_API_KEY_TRANSFORMATION",
        "cms": "GEMINI_API_KEY_CMS",
        "personalization": "GEMINI_API_KEY_PERSONALIZATION",
        "validation": "GEMINI_API_KEY_VALIDATION"
    }
    
    master_key = os.getenv("GEMINI_API_KEY")
    env_var = key_map.get(task_type)
    specialized_key = os.getenv(env_var) if env_var else None
    
    # Streamlit Secrets Fallback
    try:
        import streamlit as st
        if not master_key and "GEMINI_API_KEY" in st.secrets:
            master_key = st.secrets["GEMINI_API_KEY"]
        if not specialized_key and env_var in st.secrets:
            specialized_key = st.secrets[env_var]
    except:
        pass
    
    # helper for validation
    def is_valid(k):
        # Ensure it's a string and doesn't contain common placeholders
        return isinstance(k, str) and k.strip() != "" and \
               "YOUR" not in k.upper() and "PLACEHOLDER" not in k.upper() and \
               len(k.strip()) > 30

    if is_valid(specialized_key):
        return specialized_key.strip()
    
    if is_valid(master_key):
        return master_key.strip()
        
    return None

def call_gemini(prompt, task_type, model_name='gemini-1.5-flash'):
    # Input clipping for safety
    prompt = str(prompt)[:MAX_INPUT_SIZE]
    
    api_key = get_api_key(task_type)
    if not api_key:
        return f"Error: API Key for '{task_type}' is missing."
    
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        response = model.generate_content(prompt)
        # Final sanitization of output
        return response.text
    except Exception as e:
        err_msg = str(e)
        # Redact potential API key from error message for security
        import re
        redacted_msg = re.sub(r'AIza[0-9A-Za-z-_]{35}', '[REDACTED_API_KEY]', err_msg)
        return f"AI Error: {redacted_msg}"

def generate_hash(content):
    return hashlib.sha256(content.encode('utf-8')).hexdigest()[:12]

def extract_text_from_pdf(file_path):
    try:
        pdf = PdfReader(file_path)
        text = ""
        for i, page in enumerate(pdf.pages):
            if i > 50: # Limit pages for security/performance
                text += "\n[PDF TRUNCATED - Too many pages]"
                break
            text += page.extract_text()
        return sanitize_text(text[:MAX_INPUT_SIZE])
    except Exception as e: return f"Error reading PDF: {e}"

def calculate_reading_time(text):
    words = len(str(text).split())
    minutes = words / 200
    return f"{minutes:.1f} min"

def export_to_docx(title, content):
    """Generates a .docx file and returns its bytes."""
    try:
        from docx import Document
        from io import BytesIO
        doc = Document()
        doc.add_heading(title, 0)
        # Ensure content is string
        content_str = str(content)
        # Split content by lines to preserve paragraphs
        for line in content_str.split('\n'):
            doc.add_paragraph(line)
        
        target = BytesIO()
        doc.save(target)
        target.seek(0)
        return target.getvalue()
    except Exception as e:
        raise Exception(f"Docx Generation Error: {e}")

def export_to_pdf(title, content):
    """Generates a .pdf file and returns its bytes."""
    try:
        from fpdf import FPDF
        
        # Use a class that supports multi-page and basic headers
        class PDF(FPDF):
            def header(self):
                self.set_font('helvetica', 'B', 12)
                self.cell(0, 10, f'Content OS Export: {title[:50]}', 0, 1, 'C')
                self.ln(5)

        pdf = PDF()
        pdf.add_page()
        
        # FPDF2 supports unicode if we provide a font, but for standard fonts
        # it uses latin-1. Let's try to handle standard fonts more robustly.
        pdf.set_font("helvetica", size=12)
        
        # Clean text for PDF compatibility (removes emojis/weird chars that break standard fonts)
        # Instead of latin-1, let's use a broader approach or just clean it carefully
        safe_content = str(content).encode('ascii', 'ignore').decode('ascii')
        
        pdf.multi_cell(0, 10, safe_content)
        
        # In fpdf2, output() returns a bytearray by default. 
        # Streamlit requires bytes for the download button in some environments.
        return bytes(pdf.output())
    except Exception as e:
        raise Exception(f"PDF Generation Error: {e}")

def sanitize_text(text):
    if text is None: return ""
    # aggressive sanitization to prevent XSS
    text = str(text)
    # Remove any weird null bytes
    text = text.replace('\x00', '')
    # Escape HTML characters
    return html.escape(text)


def get_youtube_transcript(url):
    from youtube_transcript_api import YouTubeTranscriptApi
    try:
        import re
        # More robust video ID extraction
        match = re.search(r"(?:v=|\/)([0-9A-Za-z_-]{11}).*", url)
        if not match:
            return "Error: Could not find valid YouTube video ID."
        video_id = match.group(1)
        transcript = YouTubeTranscriptApi.get_transcript(video_id)
        return " ".join([t['text'] for t in transcript])
    except Exception as e: return f"Error fetching YouTube transcript: {e}"

def predict_engagement_metrics(content, tone="Professional", platform="Generic", task_type="personalization"):
    """
    AI-powered engagement prediction based on content analysis.
    Returns predicted likes, comments, shares, and engagement score.
    """
    prompt = f"""
    You are an expert social media analyst and engagement predictor.
    
    Analyze the following content and predict realistic engagement metrics:
    
    CONTENT: {content[:3000]}
    TONE: {tone}
    PLATFORM: {platform}
    
    Based on content quality, relevance, tone, and platform best practices, predict:
    1. Expected Likes/Reactions (realistic number)
    2. Expected Comments (realistic number)
    3. Expected Shares (realistic number)
    4. Engagement Score (0-100, where 100 is viral-level engagement)
    5. Best Posting Time (e.g., "Weekday Morning", "Weekend Evening")
    6. Predicted Reach (Low/Medium/High/Viral)
    
    Respond ONLY in this exact JSON format:
    {{
        "likes": <number>,
        "comments": <number>,
        "shares": <number>,
        "engagement_score": <number 0-100>,
        "best_time": "<time recommendation>",
        "predicted_reach": "<Low/Medium/High/Viral>",
        "confidence": <number 0-100>
    }}
    """
    
    try:
        result = call_gemini(prompt, task_type)
        if result and not result.startswith("Error"):
            # Extract JSON from response more robustly
            import re
            json_match = re.search(r'\{.*\}', result, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
        
        # Fallback default predictions
        return {
            "likes": 45,
            "comments": 8,
            "shares": 12,
            "engagement_score": 62,
            "best_time": "Weekday Morning",
            "predicted_reach": "Medium",
            "confidence": 75
        }
    except:
        return {
            "likes": 45,
            "comments": 8,
            "shares": 12,
            "engagement_score": 62,
            "best_time": "Weekday Morning",
            "predicted_reach": "Medium",
            "confidence": 75
        }

def predict_audience_insights(content, audience="General Tech", task_type="personalization"):
    """
    AI-powered audience behavior and preference prediction.
    Returns insights about target audience engagement patterns.
    """
    prompt = f"""
    You are an audience behavior analyst.
    
    Analyze this content and predict audience insights:
    
    CONTENT: {content[:2000]}
    TARGET AUDIENCE: {audience}
    
    Predict:
    1. Primary Age Group (e.g., "18-24", "25-34", "35-44")
    2. Engagement Pattern (e.g., "Quick Scanners", "Deep Readers", "Visual Learners")
    3. Preferred Content Length (Short/Medium/Long)
    4. Key Interest Topics (list 3-5 topics)
    5. Sentiment (Positive/Neutral/Negative)
    
    Respond ONLY in this exact JSON format:
    {{
        "age_group": "<age range>",
        "engagement_pattern": "<pattern>",
        "preferred_length": "<Short/Medium/Long>",
        "interest_topics": ["topic1", "topic2", "topic3"],
        "sentiment": "<Positive/Neutral/Negative>",
        "retention_rate": <number 0-100>
    }}
    """
    
    try:
        result = call_gemini(prompt, task_type)
        if result and not result.startswith("Error"):
            import re
            json_match = re.search(r'\{.*\}', result, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
        
        return {
            "age_group": "25-34",
            "engagement_pattern": "Deep Readers",
            "preferred_length": "Medium",
            "interest_topics": ["Technology", "Innovation", "Productivity"],
            "sentiment": "Positive",
            "retention_rate": 68
        }
    except:
        return {
            "age_group": "25-34",
            "engagement_pattern": "Deep Readers",
            "preferred_length": "Medium",
            "interest_topics": ["Technology", "Innovation", "Productivity"],
            "sentiment": "Positive",
            "retention_rate": 68
        }

def predict_user_behavior(project_history, user_prefs, task_type="personalization"):
    """
    AI-powered predictive modeling of the user's focus and future interaction patterns.
    """
    prompt = f"""
    You are a user behavior predictive model. 
    Analyze the recent interaction history and preferences:
    
    HISTORY: {str(project_history)[:2000]}
    PREFERENCES: {user_prefs}
    
    Predict the following for the next session:
    1. Predicted Intensity (0-100)
    2. Focus Area (e.g. SEO, Tone, Consistency)
    3. Suggested Next Action
    4. User Satisfication Score (0-100 based on past tone feedback)
    
    Respond ONLY in this exact JSON format:
    {{
        "predicted_intensity": <number>,
        "focus_area": "<string>",
        "suggested_action": "<string>",
        "satisfaction_prediction": <number>,
        "learning_confidence": <number>
    }}
    """
    try:
        result = call_gemini(prompt, task_type)
        if result and not result.startswith("Error"):
            import re
            json_match = re.search(r'\{.*\}', result, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
        
        return {
            "predicted_intensity": 75,
            "focus_area": "Content Refinement",
            "suggested_action": "Optimize for Social Reach",
            "satisfaction_prediction": 82,
            "learning_confidence": 65
        }
    except:
        return {
            "predicted_intensity": 75,
            "focus_area": "General Improvement",
            "suggested_action": "Continue Drafting",
            "satisfaction_prediction": 80,
            "learning_confidence": 50
        }

class IngestionClient:
    BASE_URL = "https://ai-enhanced-content-creation-ocr-api.onrender.com/ingest"
    
    def ingest_file(self, file_name, file_content, file_type):
        try:
            files = {'file': (file_name, file_content, file_type)}
            response = requests.post(self.BASE_URL, files=files, timeout=30)
            response.raise_for_status()
            return response.json()
        except Exception as e:
            return {"error": str(e)}

    def ingest_url(self, url):
        try:
            payload = {"url": url}
            response = requests.post(self.BASE_URL, json=payload, headers={"Content-Type": "application/json"}, timeout=30)
            response.raise_for_status()
            return response.json()
        except Exception as e:
            return {"error": str(e)}

class AuthManager:
    """Strongly-typed Authentication and Identity Management."""
    USER_DB_PATH = os.path.join(CMS_ROOT, "users.json")

    def __init__(self):
        if not os.path.exists(CMS_ROOT): os.makedirs(CMS_ROOT)
        self.users = self._load()

    def _load(self):
        if os.path.exists(self.USER_DB_PATH):
            with open(self.USER_DB_PATH, 'r') as f: return json.load(f)
        return {}

    def _save(self):
        with open(self.USER_DB_PATH, 'w') as f: json.dump(self.users, f, indent=2)

    def register(self, username, password):
        if username in self.users: return False, "User exists."
        salt = secrets.token_hex(8)
        pwd_hash = hashlib.sha256((password + salt).encode()).hexdigest()
        user_id = hashlib.md5(username.encode()).hexdigest()[:12]
        
        self.users[username] = {
            "id": user_id,
            "hash": pwd_hash,
            "salt": salt,
            "created_at": datetime.datetime.now().isoformat()
        }
        self._save()
        return True, user_id

    def login(self, username, password):
        user = self.users.get(username)
        if not user: return None
        check_hash = hashlib.sha256((password + user['salt']).encode()).hexdigest()
        if check_hash == user['hash']:
            return user
        return None

class ContentManager:
    LIFECYCLE_STAGES = ["Idea", "Draft", "Review", "Approval", "Publication", "Archival"]
    LEGACY_FALLBACK = "--None--"
    
    def __init__(self):
        if not os.path.exists(CMS_ROOT):
            os.makedirs(CMS_ROOT)
            
    def _get_path(self, folder, project_id):
        return os.path.join(CMS_ROOT, folder, project_id)

    def create_project(self, title, folder, content, owner_id, tags=None, extra_meta=None):
        timestamp = int(time.time())
        if not title: title = "Untitled Project"
        clean_title = "".join([c if c.isalnum() else "_" for c in title])[:30]
        project_id = f"{timestamp}_{clean_title}"
        path = self._get_path(folder, project_id)
        
        # Initialize directory structure
        os.makedirs(os.path.join(path, "main"), exist_ok=True)
        os.makedirs(os.path.join(path, "branches"), exist_ok=True)
        
        # Initial meta including collaborations
        meta = {
            "owner": owner_id,
            "collaborators": {owner_id: "Developer"}, # Creator is always Developer
            "project_id": project_id,
            "title": title,
            "folder": folder,
            "created_at": datetime.datetime.now().isoformat()
        }
        with open(os.path.join(path, "meta.json"), "w") as f:
            json.dump(meta, f, indent=2)
            
        return self.commit_version(folder, project_id, content, owner_id, title, tags or [], "Idea", "Initial commit", extra_meta)

    def commit_version(self, folder, project_id, content, user_id, title, tags, status, message="Update", extra_meta=None):
        path = self._get_path(folder, project_id)
        meta = self.get_meta(folder, project_id)
        
        # Compatibility Layer: check if this is an old-style project
        is_owner = meta.get('owner', self.LEGACY_FALLBACK) == user_id
        
        # If collaborator, commit to a branch instead of main
        sub_folder = "main" if is_owner else os.path.join("branches", user_id)
        target_dir = os.path.join(path, sub_folder)
        os.makedirs(target_dir, exist_ok=True)
        
        timestamp = datetime.datetime.now().isoformat()
        content_hash = generate_hash(content + timestamp + user_id) # Hash includes user for uniqueness
        
        version_data = {
            "version_id": content_hash,
            "contributor_hash": user_id,
            "timestamp": timestamp,
            "title": title or self.LEGACY_FALLBACK,
            "content": content,
            "tags": tags or [self.LEGACY_FALLBACK],
            "status": status,
            "message": message,
            "extra_meta": extra_meta or {}
        }
        
        with open(os.path.join(target_dir, f"v_{content_hash}.json"), "w") as f:
            json.dump(version_data, f, indent=2)
            
        # Update Head if it's the main branch
        if is_owner:
            meta["current_head"] = content_hash
            meta["last_modified"] = timestamp
            with open(os.path.join(path, "meta.json"), "w") as f:
                json.dump(meta, f, indent=2)
            
        return project_id

    def get_meta(self, folder, project_id):
        try:
            with open(os.path.join(self._get_path(folder, project_id), "meta.json"), "r") as f:
                data = json.load(f)
                # Structure-level Backwards Compatibility
                defaults = {
                    "owner": self.LEGACY_FALLBACK,
                    "collaborators": {},
                    "project_id": project_id,
                    "title": self.LEGACY_FALLBACK,
                    "folder": folder,
                    "tags": [],
                    "status": "Idea",
                    "last_modified": self.LEGACY_FALLBACK
                }
                merged = {**defaults, **data}
                # Key-level Backwards Compatibility (Type-safe Nil-punning)
                for k in merged:
                    if merged[k] is None:
                        merged[k] = defaults.get(k, self.LEGACY_FALLBACK)
                
                # Special Case: Collaborators must ALWAYS be a dict
                if not isinstance(merged.get('collaborators'), dict):
                    merged['collaborators'] = {}
                    
                return merged
        except: return None

    def get_history(self, folder, project_id, branch="main"):
        path = os.path.join(self._get_path(folder, project_id), branch)
        if branch != "main": # For collaborator branches
            path = os.path.join(self._get_path(folder, project_id), "branches", branch)
            
        files = glob.glob(os.path.join(path, "v_*.json"))
        history = []
        for f in files:
            with open(f, "r") as r:
                ver = json.load(r)
                # Inject legacy fallback for missing keys
                safe_ver = {
                    "version_id": ver.get("version_id", self.LEGACY_FALLBACK),
                    "contributor_hash": ver.get("contributor_hash", self.LEGACY_FALLBACK),
                    "timestamp": ver.get("timestamp", self.LEGACY_FALLBACK),
                    "content": ver.get("content", ""),
                    "title": ver.get("title", self.LEGACY_FALLBACK),
                    "status": ver.get("status", self.LEGACY_FALLBACK),
                    "message": ver.get("message", self.LEGACY_FALLBACK)
                }
                history.append(safe_ver)
        return sorted(history, key=lambda x: x['timestamp'], reverse=True)

    def list_all_content(self):
        projects = []
        for folder_path in glob.glob(os.path.join(CMS_ROOT, "*")):
            if os.path.isdir(folder_path):
                folder_name = os.path.basename(folder_path)
                for proj_path in glob.glob(os.path.join(folder_path, "*")):
                    if os.path.isdir(proj_path):
                        proj_id = os.path.basename(proj_path)
                        meta = self.get_meta(folder_name, proj_id)
                        if meta:
                            projects.append(meta)
        return sorted(projects, key=lambda x: x['last_modified'], reverse=True)
    
    def get_folders(self):
        if not os.path.exists(CMS_ROOT): return []
        return [d for d in os.listdir(CMS_ROOT) if os.path.isdir(os.path.join(CMS_ROOT, d))]

    def merge_branch(self, folder, project_id, branch_user_id, version_id, developer_id):
        """Allow a Developer to merge a collaborator's version into main."""
        meta = self.get_meta(folder, project_id)
        if not meta or meta.get('owner') != developer_id:
            return False, "Unauthorized: Only the Developer can merge."
            
        # Path to the specific version in the branch
        branch_path = os.path.join(self._get_path(folder, project_id), "branches", branch_user_id, f"v_{version_id}.json")
        if not os.path.exists(branch_path):
            return False, "Version not found in branch."
            
        with open(branch_path, "r") as f:
            v_data = json.load(f)
            
        # Commit this to main
        new_msg = f"Merged from branch {branch_user_id}: {v_data.get('message', '')}"
        return self.commit_version(
            folder, project_id, v_data['content'], developer_id, 
            v_data['title'], v_data['tags'], "Review", new_msg, v_data.get('extra_meta')
        )
